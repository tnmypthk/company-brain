"""
Local file ingestion: PDF and DOCX → plain text → chunks.

Why extract to plain text first? Every downstream step (chunking, embedding,
LLM prompts) works on strings. Converting once here keeps the rest of the
pipeline format-agnostic.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import List

from ingestion.chunker import Chunk, chunk_text


def _extract_pdf(path: Path) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    # Paragraphs captures body text; tables are skipped for now (Week 2 enhancement)
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def ingest_file(path: str | Path, chunk_size: int = 500, overlap: int = 50) -> List[Chunk]:
    """
    Parse a PDF or DOCX file and return a list of overlapping text chunks.

    Raises ValueError for unsupported file types so the caller gets a clear
    error instead of silent empty output.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Only .pdf and .docx are supported.")

    if not text.strip():
        return []

    return chunk_text(text, source=str(path), chunk_size=chunk_size, overlap=overlap)


def ingest_file_bytes(content: bytes, filename: str, chunk_size: int = 500, overlap: int = 50) -> List[Chunk]:
    """
    Same as ingest_file but accepts raw bytes — useful for Streamlit's file_uploader
    which gives you bytes, not a path.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() for page in reader.pages if page.extract_text()]
        text = "\n\n".join(pages)
    elif suffix == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(content))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    if not text.strip():
        return []

    return chunk_text(text, source=filename, chunk_size=chunk_size, overlap=overlap)
